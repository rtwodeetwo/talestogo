"""
Word export for the How Tales Works methodology page.

Tales does not generate written reports; this renders one static explanatory
document. The report-rendering paths that used to live here were removed with
the reports feature.
"""

import io
import re
from typing import Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy.orm import Session


def export_to_word(markdown_content: str, title: str) -> io.BytesIO:
    """
    Convert markdown report to Word document with proper formatting.

    Args:
        markdown_content: The markdown content to convert
        title: The report title

    Returns:
        BytesIO object containing the Word document
    """
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Parse markdown line by line
    lines = markdown_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # H1 Headers
        if line.startswith('# '):
            text = line[2:].strip()
            heading = doc.add_heading(text, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # H2 Headers
        elif line.startswith('## '):
            text = line[3:].strip()
            heading = doc.add_heading(text, level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # H3 Headers
        elif line.startswith('### '):
            text = line[4:].strip()
            heading = doc.add_heading(text, level=3)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Horizontal rules
        elif line.startswith('---') or line.startswith('***'):
            doc.add_paragraph('_' * 80)

        # Tables
        elif line.startswith('|'):
            # Collect all table lines
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            i -= 1  # Back up one since we'll increment at the end

            # Parse table
            if len(table_lines) >= 2:  # Header + separator minimum
                # Parse header
                header = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]

                # Skip separator line
                # Parse data rows (skip separator)
                data_rows = []
                for row_line in table_lines[2:]:  # Skip header and separator
                    cells = [cell.strip() for cell in row_line.split('|')[1:-1]]
                    data_rows.append(cells)

                # Create Word table
                if data_rows:
                    table = doc.add_table(rows=1 + len(data_rows), cols=len(header))
                    table.style = 'Light Grid Accent 1'

                    # Add header
                    header_cells = table.rows[0].cells
                    for idx, header_text in enumerate(header):
                        header_cells[idx].text = header_text
                        # Make header bold
                        for paragraph in header_cells[idx].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True

                    # Add data rows
                    for row_idx, row_data in enumerate(data_rows, start=1):
                        cells = table.rows[row_idx].cells
                        for col_idx, cell_data in enumerate(row_data):
                            # Remove markdown formatting from cell content
                            clean_text = cell_data.replace('**', '')
                            cells[col_idx].text = clean_text

        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            # Remove markdown formatting
            text = text.replace('**', '')
            doc.add_paragraph(text, style='List Bullet')

        # Numbered lists
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s+', '', line)
            # Remove markdown formatting
            text = text.replace('**', '')
            doc.add_paragraph(text, style='List Number')

        # Regular paragraphs
        else:
            # Remove markdown formatting
            text = line.replace('**', '')
            if text:
                para = doc.add_paragraph(text)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        i += 1

    # Save to BytesIO
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    return output
